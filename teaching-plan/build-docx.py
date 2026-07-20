#!/usr/bin/env python
"""
Generate Word docs from the teaching-plan JSON files.

Usage:  python build-docx.py <slug>        e.g.  python build-docx.py maths-xi

Reads   <slug>-spiral-plan.json  and  <slug>-deep-dive.json  (in this folder)
Writes  <slug>-teaching-plan.docx      (ONE combined file: plan portrait -> mapping landscape)

(build_spiral / build_deepdive remain if a standalone file is ever wanted,
 but main() emits only the combined file.)

The JSON is the source of truth; the .docx files are a throwaway rendering —
re-run this any time the JSON changes. Same generator works for every
grade/subject: just point it at a same-shaped pair of JSON files.
"""
import sys, os, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
GREY = "E7E6E6"
BRAND = RGBColor(0x37, 0x30, 0x8A)  # indigo accent for headings


# ---------- low-level helpers ----------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, size=8, bold=False, italic=False):
    """Write multi-line text into a cell (one paragraph per line)."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.text = ""
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic


def header_row(table, labels, size=8):
    for c, lab in zip(table.rows[0].cells, labels):
        set_cell(c, lab, size=size, bold=True)
        shade(c, GREY)


def set_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for c, w in zip(row.cells, widths_in):
            c.width = Inches(w)


def portrait(s):
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.7)


def landscape(s):
    if s.page_width < s.page_height:
        s.page_width, s.page_height = s.page_height, s.page_width
    s.orientation = WD_ORIENT.LANDSCAPE
    s.top_margin = s.bottom_margin = Inches(0.5)
    s.left_margin = s.right_margin = Inches(0.5)


# ---------- cell formatters ----------
def concepts_cell(concepts):
    return "\n".join("• " + c for c in concepts)


CLASS_LABEL = {"9th": "Std 9th", "10th": "Std 10th", "11th": "Std 11th", "12th": "Std 12th"}


def book_cell(d):
    """Render a textbook column (State Board / NCERT).

    New format: `homework` = [{class, ref}] renders a class-tagged homework block
      (foundation exercises from 9th/10th plus the current 11th/12th exercise).
    Legacy: a plain `exercise` string still renders as a single "→ …" line
      (kept so the other subjects' JSONs build unchanged).
    """
    sec, note = d.get("section"), d.get("note")
    hw, ex = d.get("homework"), d.get("exercise")
    if not sec and not hw:
        return "GAP" + (f"\n{note}" if note else "")
    out = []
    if sec: out.append(sec)
    if hw:
        out.append("Homework")
        for item in hw:
            cls = CLASS_LABEL.get(item.get("class", ""), item.get("class", ""))
            ref = item.get("ref", "")
            out.append("→ " + (f"{cls}: {ref}" if cls else ref))
    elif ex:
        out.append("→ " + ex)
    if note: out.append("Note: " + note)
    return "\n".join(out)


def exam_cell(d):
    cnt, topic, rel, note = d.get("pyq_count"), d.get("topic"), d.get("relevance", ""), d.get("note")
    out = []
    if cnt is not None:
        out.append(f"PYQ freq: {cnt}" + (f" ({topic})" if topic else ""))
    elif topic:
        out.append(topic)
    out.append(f"relevance: {rel}")
    if note: out.append(note)
    return "\n".join(out)


FLAG_LABEL = {
    "sb_split": "⚠ SB split", "ncert_gap": "⚠ NCERT gap",
    "ncert_aligned": "✓ NCERT", "cet_anchor": "★ CET", "nda_anchor": "★ NDA",
    "cbse_gap": "⚠ CBSE gap", "cbse_xi": "◆ CBSE: teach in XI",
    "prereq": "↑ Prerequisite (revise first)",
}


def subtopic_cell(row):
    head = f"{row['id']}  {row['subtopic']}"
    flags = [FLAG_LABEL.get(f, f) for f in row.get("flags", [])]
    return head + ("\n[" + " · ".join(flags) + "]" if flags else "")


# ---------- content bodies (add to a given doc; no create/save) ----------
def spine_body(doc, data):
    title = data.get("plan_title", "Spiral Teaching Plan")
    prefix = data.get("title_prefix", "LWS Pune ")  # set "" to de-brand the main title
    h = doc.add_heading(f"{prefix}{data['subject']} Std {roman(data['grade'])} — {title}", level=0)
    h.runs[0].font.color.rgb = BRAND
    meta = doc.add_paragraph()
    meta.add_run(f"{data['board']} · {data['textbook']}").italic = True

    # Overview — the operational roadmap. Timeline + Sign are blank fill-in columns:
    # the teacher writes planned dates in Timeline and signs Sign as each unit completes.
    # (Phase per unit still shows in the mapping section; sessions live in the JSON.)
    doc.add_heading("Chapter Overview — All Units", level=1)
    units = data["units"]
    t = doc.add_table(rows=len(units) + 1, cols=5); t.style = "Table Grid"
    header_row(t, ["#", "Chapter Title", "Board Source", "Timeline", "Sign"], size=9)
    for i, u in enumerate(units, start=1):
        set_cell(t.rows[i].cells[0], u["unit_no"], size=9, bold=True)
        set_cell(t.rows[i].cells[1], u["title"], size=9)
        set_cell(t.rows[i].cells[2], u["board_source"], size=9)
        set_cell(t.rows[i].cells[3], "", size=9)   # Timeline — teacher fills in planned dates
        set_cell(t.rows[i].cells[4], "", size=9)   # Sign — completion signature
    set_widths(t, [0.5, 2.7, 1.9, 1.2, 0.9])

    supp = data.get("supplementary_nda_topics")
    if supp:
        doc.add_heading("Supplementary NDA Topics — not in the board textbook (to fit later)", level=1)
        for s in supp:
            para = doc.add_paragraph()
            para.add_run(f"{s['topic']} ").bold = True
            para.add_run(f"[{s['status']}] — {s['note']}").font.size = Pt(9)


def deepdive_body(doc, data, unit_meta=None, first_break=False):
    unit_meta = unit_meta or {}
    h = doc.add_heading(f"LWS Pune {data['subject']} Std {roman(data['grade'])} — Syllabus Mapping", level=0)
    h.runs[0].font.color.rgb = BRAND
    doc.add_paragraph().add_run(
        "State Board / CET / NCERT / NDA · ★ = primary exam focus · "
        "✓ = clean cross-book alignment · ⚠ = split or gap to flag · "
        "PYQ freq = past-year-question count for the topic cluster (do not sum across rows) · "
        "relevance = editorial exam-emphasis").italic = True

    cols = ["Subtopic", "Concepts", "State Board", "CET", "NCERT", "NDA"]
    widths = [1.5, 2.5, 1.6, 1.5, 1.6, 1.5]
    for ui, u in enumerate(data["units"]):
        m = unit_meta.get(u["unit_no"], {})
        title = m.get("title", "")
        head = doc.add_heading(f"Unit {u['unit_no']}" + (f" · {title}" if title else ""), level=1)
        if ui > 0 or first_break:
            head.paragraph_format.page_break_before = True  # each unit on its own page
        bits = [b for b in (m.get("board_source"),
                            f"Phase {m['phase']}" if m.get("phase") else None,
                            f"{m['sessions']} sessions" if m.get("sessions") else None) if b]
        if bits:
            doc.add_paragraph().add_run(" · ".join(bits)).italic = True
        t = doc.add_table(rows=len(u["rows"]) + 1, cols=6); t.style = "Table Grid"
        header_row(t, cols)
        for i, row in enumerate(u["rows"], start=1):
            cells = t.rows[i].cells
            set_cell(cells[0], subtopic_cell(row), bold=False)
            set_cell(cells[1], concepts_cell(row["concepts"]))
            set_cell(cells[2], book_cell(row["state_board"]))
            set_cell(cells[3], exam_cell(row["cet"]))
            set_cell(cells[4], book_cell(row["ncert"]))
            set_cell(cells[5], exam_cell(row["nda"]))
        set_widths(t, widths)


def principles_body(doc, data):
    """Guiding Principles table — the design rationale, rendered at the END of the plan."""
    doc.add_heading("Guiding Principles", level=1)
    gp = data["guiding_principles"]
    t = doc.add_table(rows=len(gp) + 1, cols=2); t.style = "Table Grid"
    header_row(t, ["Principle", "What it means"], size=9)
    for i, g in enumerate(gp, start=1):
        set_cell(t.rows[i].cells[0], g["principle"], size=9, bold=True)
        set_cell(t.rows[i].cells[1], g["meaning"], size=9)
    set_widths(t, [1.6, 5.6])


# ---------- document builders ----------
def new_doc():
    doc = Document(); doc.styles["Normal"].font.name = "Calibri"
    return doc


def build_spiral(data, out_path):
    doc = new_doc(); portrait(doc.sections[0])
    spine_body(doc, data)
    principles_body(doc, data)
    doc.save(out_path); return out_path


def build_deepdive(data, out_path, unit_meta=None):
    doc = new_doc(); landscape(doc.sections[0])
    deepdive_body(doc, data, unit_meta)
    doc.save(out_path); return out_path


def build_combined(spine, deep, unit_meta, out_path):
    doc = new_doc(); portrait(doc.sections[0])
    spine_body(doc, spine)                            # title + Chapter Overview (+ supplementary)
    landscape(doc.add_section(WD_SECTION.NEW_PAGE))   # section break: overview -> mapping
    deepdive_body(doc, deep, unit_meta)
    portrait(doc.add_section(WD_SECTION.NEW_PAGE))    # section break: mapping -> guiding principles (at end)
    principles_body(doc, spine)
    doc.save(out_path); return out_path


def roman(n):
    return {9: "IX", 10: "X", 11: "XI", 12: "XII"}.get(n, str(n))


def main():
    if len(sys.argv) < 2:
        sys.exit("Usage: python build-docx.py <slug>   (e.g. maths-xi)")
    slug = sys.argv[1]
    spine_p = os.path.join(HERE, f"{slug}-spiral-plan.json")
    deep_p = os.path.join(HERE, f"{slug}-deep-dive.json")
    spine = json.load(open(spine_p, encoding="utf-8")) if os.path.exists(spine_p) else None
    deep = json.load(open(deep_p, encoding="utf-8")) if os.path.exists(deep_p) else None
    if not (spine and deep):
        sys.exit(f"Need both {slug}-spiral-plan.json and {slug}-deep-dive.json in {HERE}")
    unit_meta = {u["unit_no"]: u for u in spine["units"]}
    out = build_combined(spine, deep, unit_meta, os.path.join(HERE, f"{slug}-teaching-plan.docx"))
    print("wrote:", os.path.basename(out))


if __name__ == "__main__":
    main()
